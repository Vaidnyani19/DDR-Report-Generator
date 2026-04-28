from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os

# UrbanRoof brand colors
ORANGE = RGBColor(0xE8, 0x7A, 0x1E)    # UrbanRoof orange
DARK_GRAY = RGBColor(0x40, 0x40, 0x40)  # Dark gray text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)       # White
RED = RGBColor(0xC0, 0x00, 0x00)        # High severity
AMBER = RGBColor(0xFF, 0x8C, 0x00)      # Medium severity
GREEN = RGBColor(0x37, 0x86, 0x24)      # Low severity
LIGHT_ORANGE_BG = "FFF3E0"              # Light orange background
HEADER_BG = "E87A1E"                    # Orange header background

def set_cell_background(cell, hex_color):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_horizontal_line(doc, color="E87A1E"):
    """Add a colored horizontal line"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_section_heading(doc, number, title):
    """Add orange section heading with number"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ORANGE
    add_horizontal_line(doc)
    return p

def add_label_value(doc, label, value, label_bold=True):
    """Add a label: value paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = label_bold
    label_run.font.size = Pt(11)
    label_run.font.color.rgb = DARK_GRAY
    value_run = p.add_run(str(value))
    value_run.font.size = Pt(11)
    return p

def find_best_image(hint, image_list, preferred_source=None, used_paths=None):
    """
    Find the best matching image for an area.
    Args:
        hint: keyword like 'hall', 'bedroom', 'bathroom'
        image_list: full list of extracted images
        preferred_source: 'inspection' or 'thermal'
        used_paths: set of already used image paths to avoid duplicates
    """
    if used_paths is None:
        used_paths = set()
    
    hint_lower = hint.lower()
    candidates = image_list
    
    # First try preferred source
    if preferred_source:
        candidates = [img for img in image_list if img["source"] == preferred_source]
    
    # Try keyword match in filename
    for img in candidates:
        if img["path"] not in used_paths:
            filename = os.path.basename(img["path"]).lower()
            if hint_lower in filename:
                used_paths.add(img["path"])
                return img["path"]
    
    # Fallback: return any unused image from preferred source
    for img in candidates:
        if img["path"] not in used_paths:
            used_paths.add(img["path"])
            return img["path"]
    
    # Last resort: any unused image
    for img in image_list:
        if img["path"] not in used_paths:
            used_paths.add(img["path"])
            return img["path"]
    
    return None

def insert_image_safe(doc, img_path, width_inches=4.0, caption=None):
    """Safely insert image with error handling"""
    if img_path and os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(width_inches))
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].italic = True
        except Exception as e:
            print(f"Could not insert image {img_path}: {e}")
            doc.add_paragraph(f"[Image could not be loaded]")
    else:
        doc.add_paragraph("[Image Not Available]")

def verify_images(image_list):
    print(f"Total images passed to docx builder: {len(image_list)}")
    for img in image_list[:5]:
        path = img['path']
        exists = os.path.exists(path)
        size = os.path.getsize(path) if exists else 0
        print(f"  {path} | exists={exists} | size={size} bytes")

def build_docx(ddr_data, all_images, thermal_maps, real_photos, output_path):
    verify_images(all_images)
    """
    Build professional DDR Word document.
    Args:
        ddr_data: parsed JSON from Groq
        all_images: all extracted images (inspection + thermal)
        thermal_maps: thermal heat map images (first image per thermal page)
        real_photos: real photos from thermal PDF (second image per thermal page)
        output_path: where to save the .docx
    """
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    
    used_inspection_paths = set()
    used_thermal_paths = set()
    
    # Separate inspection images
    inspection_images = [img for img in all_images if img["source"] == "inspection"]
    
    # ==========================================
    # COVER SECTION
    # ==========================================
    
    # Orange header bar using table
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.rows[0].cells[0]
    set_cell_background(header_cell, "E87A1E")
    header_cell.width = Inches(6.5)
    
    title_para = header_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(20)
    title_para.paragraph_format.space_after = Pt(20)
    title_run = title_para.add_run("Detailed Diagnostic Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = WHITE
    title_run.font.name = "Arial"
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph("Prepared by UrbanRoof Inspection System")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = DARK_GRAY
    
    # Property info table
    prop_info = ddr_data.get("property_info", {})
    if any(prop_info.values()):
        doc.add_paragraph()
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Table Grid"
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_rows = [
            ("Property Type", prop_info.get("property_type", "Not Available")),
            ("Inspection Date", prop_info.get("inspection_date", "Not Available")),
            ("Inspected By", prop_info.get("inspected_by", "Not Available")),
            ("Inspection Score", prop_info.get("inspection_score", "Not Available")),
        ]
        
        for i, (label, value) in enumerate(info_rows):
            row = info_table.rows[i]
            set_cell_background(row.cells[0], "FFF3E0")
            label_p = row.cells[0].paragraphs[0]
            label_run = label_p.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(10)
            
            val_p = row.cells[1].paragraphs[0]
            val_p.add_run(str(value)).font.size = Pt(10)
    
    doc.add_page_break()
    
    # ==========================================
    # SECTION 1: PROPERTY ISSUE SUMMARY
    # ==========================================
    add_section_heading(doc, 1, "Property Issue Summary")
    summary_p = doc.add_paragraph(ddr_data.get("property_issue_summary", "Not Available"))
    summary_p.paragraph_format.space_after = Pt(12)
    
    # ==========================================
    # SECTION 2: AREA-WISE OBSERVATIONS
    # ==========================================
    add_section_heading(doc, 2, "Area-wise Observations")
    
    areas = ddr_data.get("area_wise_observations", [])
    
    for i, area_data in enumerate(areas):
        area_name = area_data.get("area", f"Area {i+1}")
        hint = area_data.get("image_hint", "").lower()
        severity = area_data.get("severity", "Medium")
        
        # Area subheading with colored box
        area_table = doc.add_table(rows=1, cols=1)
        area_cell = area_table.rows[0].cells[0]
        set_cell_background(area_cell, "F5F5F5")
        area_p = area_cell.paragraphs[0]
        area_p.paragraph_format.space_before = Pt(6)
        area_p.paragraph_format.space_after = Pt(6)
        area_num_run = area_p.add_run(f"Area {i+1}: ")
        area_num_run.bold = True
        area_num_run.font.size = Pt(13)
        area_num_run.font.color.rgb = ORANGE
        area_name_run = area_p.add_run(area_name.upper())
        area_name_run.bold = True
        area_name_run.font.size = Pt(13)
        area_name_run.font.color.rgb = DARK_GRAY
        
        doc.add_paragraph()
        
        # Observations table
        obs_table = doc.add_table(rows=3, cols=2)
        obs_table.style = "Table Grid"
        
        obs_rows = [
            ("📍 Observation (Impacted Side)", area_data.get("negative_observation", "Not Available")),
            ("🔍 Source of Issue (Exposed Side)", area_data.get("positive_source", "Not Available")),
            ("🌡️ Thermal Finding", area_data.get("thermal_finding", "Not Available")),
        ]
        
        for j, (label, value) in enumerate(obs_rows):
            row = obs_table.rows[j]
            set_cell_background(row.cells[0], "FFF3E0")
            
            lp = row.cells[0].paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(10)
            
            vp = row.cells[1].paragraphs[0]
            vr = vp.add_run(str(value))
            vr.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # IMAGES - inspection photo + thermal image side by side
        # Try to get one inspection image and one thermal image
        
        insp_img = find_best_image(hint, inspection_images, 
                                    preferred_source="inspection",
                                    used_paths=used_inspection_paths)
        
        # For thermal: use page index matching area number
        thermal_img = None
        thermal_map = None
        if i < len(real_photos):
            thermal_img = real_photos[i]["path"]
            used_thermal_paths.add(thermal_img)
        if i < len(thermal_maps):
            thermal_map = thermal_maps[i]["path"]
            used_thermal_paths.add(thermal_map)
        
        # Image layout: sequential paragraphs
        if insp_img:
            print(f"Inserting image: {insp_img}, exists: {os.path.exists(insp_img)}, size: {os.path.getsize(insp_img) if os.path.exists(insp_img) else 0}")
            if os.path.exists(insp_img):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(insp_img, width=Inches(4.5))
                    print(f"SUCCESS: inserted {insp_img}")
                    
                    cap = doc.add_paragraph("Site Inspection Photo")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cap.runs:
                        cap.runs[0].font.size = Pt(9)
                        cap.runs[0].italic = True
                except Exception as e:
                    print(f"FAILED to insert {insp_img}: {type(e).__name__}: {e}")
                    doc.add_paragraph(f"[Image error: {e}]").alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("[Inspection Image Not Found]").alignment = WD_ALIGN_PARAGRAPH.CENTER

        thermal_to_use = thermal_map or thermal_img
        if thermal_to_use:
            print(f"Inserting image: {thermal_to_use}, exists: {os.path.exists(thermal_to_use)}, size: {os.path.getsize(thermal_to_use) if os.path.exists(thermal_to_use) else 0}")
            if os.path.exists(thermal_to_use):
                try:
                    p2 = doc.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run2 = p2.add_run()
                    run2.add_picture(thermal_to_use, width=Inches(4.5))
                    print(f"SUCCESS: inserted {thermal_to_use}")
                    
                    cap2 = doc.add_paragraph("Thermal Scan")
                    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cap2.runs:
                        cap2.runs[0].font.size = Pt(9)
                        cap2.runs[0].italic = True
                except Exception as e:
                    print(f"FAILED to insert {thermal_to_use}: {type(e).__name__}: {e}")
                    doc.add_paragraph(f"[Image error: {e}]").alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("[Thermal Scan Not Found]").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        # If thermal also has a real photo, add it full width below
        if thermal_img and thermal_map:
            print(f"Inserting image: {thermal_img}, exists: {os.path.exists(thermal_img)}, size: {os.path.getsize(thermal_img) if os.path.exists(thermal_img) else 0}")
            if os.path.exists(thermal_img):
                try:
                    rp = doc.add_paragraph()
                    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run3 = rp.add_run()
                    run3.add_picture(thermal_img, width=Inches(4.5))
                    print(f"SUCCESS: inserted {thermal_img}")
                    
                    cap3 = doc.add_paragraph("Thermal Reference Photo")
                    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cap3.runs:
                        cap3.runs[0].font.size = Pt(9)
                        cap3.runs[0].italic = True
                except Exception as e:
                    print(f"FAILED to insert {thermal_img}: {type(e).__name__}: {e}")
                    doc.add_paragraph(f"[Image error: {e}]").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        add_horizontal_line(doc, color="CCCCCC")
        doc.add_paragraph()
    
    # ==========================================
    # SECTION 3: SUMMARY TABLE
    # ==========================================
    add_section_heading(doc, 3, "Issue Summary Table")
    
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    
    # Header row
    headers = ["Area", "Issue (Impacted Side)", "Source (Exposed Side)"]
    for j, header in enumerate(headers):
        cell = summary_table.rows[0].cells[j]
        set_cell_background(cell, "E87A1E")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    
    for area_data in areas:
        row = summary_table.add_row()
        row.cells[0].paragraphs[0].add_run(
            area_data.get("area", "")).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(
            area_data.get("negative_observation", "")).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(
            area_data.get("positive_source", "")).font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ==========================================
    # SECTION 4: PROBABLE ROOT CAUSE
    # ==========================================
    add_section_heading(doc, 4, "Probable Root Cause")
    doc.add_paragraph(ddr_data.get("probable_root_cause", "Not Available"))
    
    # ==========================================
    # SECTION 5: SEVERITY ASSESSMENT
    # ==========================================
    add_section_heading(doc, 5, "Severity Assessment")
    
    severity_data = ddr_data.get("severity_assessment", {})
    level = severity_data.get("overall_level", "Medium").upper()
    
    sev_table = doc.add_table(rows=2, cols=2)
    sev_table.style = "Table Grid"
    
    # Severity level cell
    level_cell = sev_table.rows[0].cells[0]
    level_cell2 = sev_table.rows[0].cells[1]
    set_cell_background(level_cell, "F5F5F5")
    lp = level_cell.paragraphs[0]
    lp.add_run("Severity Level:").bold = True
    
    level_color = "C00000" if level == "HIGH" else "FF8C00" if level == "MEDIUM" else "378624"
    level_p = level_cell2.paragraphs[0]
    level_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(level_cell2, level_color)
    level_run = level_p.add_run(f"  {level}  ")
    level_run.bold = True
    level_run.font.color.rgb = WHITE
    level_run.font.size = Pt(14)
    
    # Reasoning row spanning both columns
    reason_row = sev_table.rows[1]
    reason_cell = reason_row.cells[0].merge(reason_row.cells[1])
    reason_cell.paragraphs[0].add_run(
        severity_data.get("reasoning", "Not Available")).font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ==========================================
    # SECTION 6: RECOMMENDED ACTIONS
    # ==========================================
    add_section_heading(doc, 6, "Recommended Actions")
    
    actions = ddr_data.get("recommended_actions", [])
    
    # Group by priority
    for priority_label in ["Immediate", "Short-term", "Long-term"]:
        priority_actions = [a for a in actions 
                           if isinstance(a, dict) and 
                           a.get("priority", "") == priority_label]
        if priority_actions:
            p = doc.add_paragraph()
            run = p.add_run(f"{priority_label} Actions:")
            run.bold = True
            run.font.color.rgb = ORANGE
            
            for action in priority_actions:
                action_p = doc.add_paragraph(style="List Number")
                area_run = action_p.add_run(
                    f"[{action.get('area', 'General')}] ")
                area_run.bold = True
                action_p.add_run(action.get("action", ""))
    
    # Handle case where actions are plain strings
    string_actions = [a for a in actions if isinstance(a, str)]
    for j, action in enumerate(string_actions, 1):
        action_p = doc.add_paragraph(style="List Number")
        action_p.add_run(action)
    
    doc.add_paragraph()
    
    # ==========================================
    # SECTION 7: CHECKLIST FINDINGS  
    # ==========================================
    checklist = ddr_data.get("checklist_findings", {})
    if checklist:
        add_section_heading(doc, 7, "Checklist Findings")
        
        check_table = doc.add_table(rows=len(checklist), cols=2)
        check_table.style = "Table Grid"
        
        for j, (key, value) in enumerate(checklist.items()):
            row = check_table.rows[j]
            set_cell_background(row.cells[0], "FFF3E0")
            label = key.replace("_", " ").title()
            row.cells[0].paragraphs[0].add_run(label).bold = True
            
            val_cell = row.cells[1].paragraphs[0]
            val_run = val_cell.add_run(str(value))
            if value == "Yes":
                val_run.font.color.rgb = RED
            elif value == "No":
                val_run.font.color.rgb = GREEN
        
        doc.add_paragraph()
    
    # ==========================================
    # SECTION 8: ADDITIONAL NOTES
    # ==========================================
    add_section_heading(doc, 8, "Additional Notes")
    doc.add_paragraph(ddr_data.get("additional_notes", "None"))
    
    # ==========================================
    # SECTION 9: MISSING OR UNCLEAR INFO
    # ==========================================
    add_section_heading(doc, 9, "Missing or Unclear Information")
    missing = ddr_data.get("missing_or_unclear_info", [])
    if missing and len(missing) > 0:
        for item in missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(item))
    else:
        doc.add_paragraph("None — All required information was available in the documents.")
    
    # Save
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    print(f"DDR Report saved to: {output_path}")
    return output_path
